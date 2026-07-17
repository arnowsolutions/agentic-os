"""
Workflow Runner — Daisy-chain LLM steps for meeting transcription → report pipelines.

Supports:
  - meeting-report: Audio/text → structured report with decisions, action items, attendees
  - Future: content-engine, policy-email, etc.

Output saved to <workspace>/workflow-outputs/<workflow>/<timestamp>/
"""

import json
import logging
import os
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

from modules.config import get_settings
from modules.llm_client import chat_completion

logger = logging.getLogger("agentic_os.workflow_runner")

# ─── Output Paths ─────────────────────────────────────────────────

WORKSPACE_ROOT = Path(__file__).resolve().parent.parent.parent  # /workspace
OUTPUT_BASE = WORKSPACE_ROOT / "workflow-outputs"


def _run_id() -> str:
    return datetime.now().strftime("%Y-%m-%d_%H%M%S")


def _output_dir(workflow: str) -> Path:
    d = OUTPUT_BASE / workflow / _run_id()
    d.mkdir(parents=True, exist_ok=True)
    return d


# ─── LLM Helpers ──────────────────────────────────────────────────

def _call_llm(system: str, user: str, temperature: float = 0.2) -> str:
    """Single-turn LLM call via OpenRouter, returns text content."""
    settings = get_settings()
    messages = [
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ]
    try:
        resp = chat_completion(messages)
        content = resp.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
        if content.startswith("```"):
            lines = content.split("\n")
            content = "\n".join(lines[1:])
            if content.endswith("```"):
                content = content[:-3].strip()
        return content
    except Exception as e:
        logger.error("LLM call failed: %s", e)
        raise


# ─── Stage 1: Extract Text from Files ──────────────────────────────

def _extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF, DOCX, or other document formats."""
    ext = os.path.splitext(file_path)[1].lower()

    # Plain text formats — direct read
    if ext in (".txt", ".md", ".json", ".csv", ".log"):
        return Path(file_path).read_text()

    # PDF
    if ext == ".pdf":
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            text = "\n\n".join(text_parts)
            if text.strip():
                return text
        except Exception as e:
            logger.warning("pdfplumber extraction failed: %s", e)

        # Fallback: try pymupdf
        try:
            import fitz
            doc = fitz.open(file_path)
            text = "\n\n".join(page.get_text() for page in doc)
            doc.close()
            if text.strip():
                return text
        except ImportError:
            pass

        raise RuntimeError(f"Could not extract text from PDF: {file_path}")

    # DOCX
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text.strip():
                        text += "\n" + row_text
            if text.strip():
                return text
        except Exception as e:
            logger.warning("python-docx extraction failed: %s", e)
        raise RuntimeError(f"Could not extract text from DOCX: {file_path}")

    # Audio files — try whisper
    if ext in (".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".mpga", ".mp4", ".mpeg"):
        return _transcribe_audio(file_path)

    # Unknown — try as plain text, fall back to audio
    try:
        return Path(file_path).read_text()
    except UnicodeDecodeError:
        return _transcribe_audio(file_path)


def _transcribe_audio(audio_path: str) -> str:
    """Transcribe audio file to text using whisper CLI."""
    try:
        r = subprocess.run(
            ["whisper", audio_path, "--model", "tiny", "--output_dir", "/tmp", "--output_format", "txt"],
            capture_output=True, text=True, timeout=300,
        )
        if r.returncode == 0:
            import glob
            txt_files = glob.glob("/tmp/*.txt")
            if txt_files:
                newest = max(txt_files, key=os.path.getmtime)
                return Path(newest).read_text()
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    raise RuntimeError(
        "Audio transcription requires openai-whisper (`pip install openai-whisper`). "
        "For now, paste the transcript text directly."
    )


def _transcribe_input(file_path: Optional[str] = None, text: Optional[str] = None) -> str:
    """Get transcript from file or direct text input."""
    if text and text.strip():
        return text.strip()
    if file_path and os.path.isfile(file_path):
        return _extract_text_from_file(file_path)
    raise ValueError("Provide either text or a valid file path.")


# ─── File Upload Handler ───────────────────────────────────────────

UPLOAD_DIR = WORKSPACE_ROOT / "workflow-uploads"

def save_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """Save an uploaded file and return its absolute path."""
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = f"{ts}_{filename}"
    file_path = UPLOAD_DIR / safe_name
    file_path.write_bytes(file_bytes)
    return str(file_path)


# ─── Stage 2: Extract Structured Data ─────────────────────────────

EXTRACT_SYSTEM = """You are a precise meeting analyst. Given a meeting transcript, extract structured data.

Return a JSON object with these keys:
- title: A concise, descriptive meeting title (1 line)
- date: Meeting date if mentioned, otherwise null
- attendees: Array of {name: string, role: string} for each person mentioned
- decisions: Array of {decision: string, context: string} — formal decisions made
- action_items: Array of {item: string, owner: string, deadline: string|null}
- key_discussion_points: Array of strings — 3-7 most important topics discussed
- summary: 2-4 paragraph executive summary of the entire meeting
- next_steps: Array of strings — what happens next
- tags: Array of strings — 3-5 topic tags (e.g. "scheduling", "compliance", "onboarding")

Be thorough. If something is not mentioned, use an empty array/null rather than fabricating.
Return ONLY the JSON object, no markdown, no explanation."""


def _extract_structured_data(transcript: str) -> dict:
    """Extract meeting metadata from transcript via LLM."""
    try:
        raw = _call_llm(EXTRACT_SYSTEM, f"Meeting transcript:\n\n{transcript}")
        return json.loads(raw)
    except json.JSONDecodeError:
        # Retry with stricter formatting
        retry_prompt = f"{EXTRACT_SYSTEM}\n\nCRITICAL: Your response must be valid JSON and nothing else. Do not wrap in ``` markers."
        raw = _call_llm(retry_prompt, f"Meeting transcript:\n\n{transcript[:8000]}")
        # Try to extract JSON from the response
        raw = raw.strip()
        if raw.startswith("```"):
            raw = "\n".join(raw.split("\n")[1:])
            if raw.endswith("```"):
                raw = raw[:-3].strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError as e:
            logger.error("Failed to parse extraction JSON: %s", e)
            return {
                "title": "Untitled Meeting",
                "date": None,
                "attendees": [],
                "decisions": [],
                "action_items": [],
                "key_discussion_points": [],
                "summary": raw[:1000],
                "next_steps": [],
                "tags": [],
            }


# ─── Stage 3: Generate Report ─────────────────────────────────────

REPORT_SYSTEM = """You are a professional report writer for a healthcare administration department (Montefiore Urology).
Given structured meeting data, produce a clean, professional meeting report in markdown.

The report must include:
1. Title + date header
2. Attendees table
3. Executive Summary
4. Key Discussion Points (numbered)
5. Decisions Made (table with context)
6. Action Items (table with owner + deadline)
7. Next Steps
8. Tags at bottom

Use clean formatting. No fluff. Professional tone. Use the department's style: clean, direct, actionable.

Return the complete report in markdown. Do NOT wrap in ``` — just the raw markdown."""


def _generate_report(data: dict) -> str:
    """Generate formatted meeting report from structured data."""
    prompt = json.dumps(data, indent=2)
    report = _call_llm(REPORT_SYSTEM, f"Meeting data:\n\n{prompt}", temperature=0.3)

    # Clean up any stray fences
    report = report.strip()
    if report.startswith("```"):
        lines = report.split("\n")
        report = "\n".join(lines[1:])
        if report.endswith("```"):
            report = report[:-3].strip()

    return report


# ─── Main Pipeline ─────────────────────────────────────────────────

def run_meeting_report(
    file_path: Optional[str] = None,
    text: Optional[str] = None,
    email_recipients: Optional[str] = None,
) -> dict:
    """
    Full pipeline: meeting audio/text → structured report → saved to disk.

    Returns:
        dict with keys: success, report, report_path, data, output_dir, error (if failed)
    """
    try:
        # Stage 1: Transcribe
        transcript = _transcribe_input(file_path=file_path, text=text)

        # Stage 2: Extract
        data = _extract_structured_data(transcript)

        # Stage 3: Generate Report
        report = _generate_report(data)

        # Stage 4: Save
        out_dir = _output_dir("meeting-report")
        report_path = out_dir / "report.md"

        # Full report
        report_path.write_text(report)

        # Structured data as JSON (for programmatic use)
        data_path = out_dir / "data.json"
        data_path.write_text(json.dumps(data, indent=2))

        # Raw transcript
        transcript_path = out_dir / "transcript.txt"
        transcript_path.write_text(transcript)

        logger.info("Meeting report saved to %s", out_dir)

        result = {
            "success": True,
            "report": report,
            "report_path": str(report_path),
            "data": data,
            "output_dir": str(out_dir),
            "transcript_length": len(transcript),
        }

        # Optional: email
        if email_recipients:
            try:
                _send_report_email(report, email_recipients, data.get("title", "Meeting Report"))
                result["email_sent"] = True
            except Exception as e:
                logger.warning("Email send failed: %s", e)
                result["email_sent"] = False
                result["email_error"] = str(e)

        return result

    except ValueError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        logger.exception("Meeting report pipeline failed")
        return {"success": False, "error": str(e)}


def _send_report_email(report: str, recipients: str, subject: str) -> None:
    """Send the report via email using available SMTP/API."""
    from modules.smtp_sender import send_email as smtp_send

    recipients_list = [r.strip() for r in recipients.split(",") if r.strip()]
    if not recipients_list:
        raise ValueError("No valid recipients")

    smtp_send(
        to_emails=recipients_list,
        subject=f"Meeting Report: {subject}",
        body=report,
        content_type="text/markdown",
    )


# ─── Workflow Registry ─────────────────────────────────────────────

WORKFLOWS = {
    "meeting-report": {
        "name": "Meeting → Report",
        "description": "Transcribe meeting audio/text and generate a structured report with decisions, action items, and summary.",
        "runner": run_meeting_report,
        "inputs": {
            "file_path": {"type": "string", "required": False, "description": "Path to audio or text file"},
            "text": {"type": "string", "required": False, "description": "Direct meeting text/transcript"},
            "email_recipients": {"type": "string", "required": False, "description": "Comma-separated email addresses"},
        },
    },
}


def list_workflows() -> list:
    """Return available workflow definitions."""
    return [
        {
            "id": wid,
            "name": w["name"],
            "description": w["description"],
            "inputs": w["inputs"],
        }
        for wid, w in WORKFLOWS.items()
    ]


def run_workflow(workflow_id: str, **kwargs) -> dict:
    """Run a workflow by ID."""
    if workflow_id not in WORKFLOWS:
        return {"success": False, "error": f"Unknown workflow: {workflow_id}. Available: {list(WORKFLOWS.keys())}"}
    runner = WORKFLOWS[workflow_id]["runner"]
    return runner(**kwargs)
